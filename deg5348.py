#!/usr/bin/env python
# coding: utf-8

# In[2]:


import numpy as np

# In[3]:


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CustomDocument():
  def __init__(self) -> None:
    self.content=""
    self.mathContentArray = []
    self.contentArray = []
    self.contentaction = ""
    self.doc = Document()
    

  def fullcontent(self,file_name="output.docx",font_name='Times New Roman', font_size=10):
    """_summary_ To load the full python file as docx

    Args:
        file_name (str, optional): Name of File. Defaults to "output.docx".
        font_name (str, optional): _description_. Defaults to 'Times New Roman'.
        font_size (int, optional): _description_. Defaults to 10.
    """
      # Create a new Document
      # Add each line of content to the document
    try:
      
      papra =self.doc.add_paragraph("\n".join(self.contentArray))
      if len(papra.runs) > 0:
        papra_run = papra.runs[0]
        papra_run.font.name = font_name
        papra_run.font.size = Pt(font_size) 

        # Save the document to the specified file
      self.doc.save(file_name)
      self.contentArray=[]
    except Exception as e:
      print("Exception : ",e)

  def contentmatrix(self,row,col,sumresult):
    # contentModify = "\n\nV^*_[{row}][{col}] = max({content})  = {res} \n".format(row=row+1,col=col+1,content=self.contentaction,res=sumresult)
    contentModify = "\n\nV^*_[{row}][{col}] = max({content})  = {res} \n".format(row=row+1,col=col+1,content=self.contentaction,res=sumresult)
    self.contentArray.append("----------------------------------------------------------------------------------------------")
    self.contentArray.append(contentModify)
    self.contentaction=""
    
  def addTable(self,data, heading, font_name='Times New Roman', font_size=10):
    heading_paragraph = self.doc.add_heading(level=1)
    heading_run = heading_paragraph.add_run(heading)
    heading_run.font.name = font_name
    heading_run.font.size = Pt(font_size + 2)  
    heading_run.bold = True
    
    # Add a table with the number of rows and columns based on the data
    table = self.doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'  # Optional: Apply a table style
    
    # Populate the table with data and format the text
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            paragraph = cell.add_paragraph(str(cell_text))
            run = paragraph.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = True  # Make the text bold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # self.contentArray.append(table)
    
  def enwrapactions(self,action,sumValue):
    """_summary_ To Put in the document actions and 

    Args:
        action (_type_): Actions agent takes
        sumValue (_type_): The Sum value received by Bellman Equation
    """
    print(sumValue)
    self.contentaction += "\n{act}<{cont} = {sumValue}>\n,".format(act=action,cont=self.content,sumValue=sumValue[action])
    self.content=""
     


# The equation is V<sub>k</sub> = max(up,down,left,right)
# <br>
# up=0,down=1,right=2,left=3
# <br>
# $$
# V^*_s = \ P(s'|a,s) \left[ R_{s,a,s'} + \gamma V^*_{s'} \right]
# $$
# 
# 

# In[4]:


GRID_SIZE = 4

# In[5]:


def is_valid_grid(row,col,grid_size):
    return 0 <= row < grid_size and 0 <= col < grid_size

# In[6]:



def next_position(row,col,action):
  """_summary_ Send the position the agent takes

  Args:
      row (_type_): Row of answer matrix
      col (_type_): Column of answer matrix
      action (_type_): Which action to take
      rewardmatrix (_type_): Matrix to  receive 

  Returns:
      _type_: _description_
  """
  p_row = row
  p_col = col
  if action == 0:
    p_row -= 1
    
  if action == 1:
    p_row +=1
    
  if action == 2:
    p_col +=1
    
  if action == 3:
    p_col -=1
    
  return p_row,p_col

# In[7]:


def valueIterations(GAMMA, Trans, rewardmatrix, ansmatrix, QArr, actions,iterations=1):
    
    for V in range(0,iterations):
      customdocument = CustomDocument()
      temp_ansmatrix = np.copy(ansmatrix)
      
      # Traverse row and column
      for row in range(0,len(rewardmatrix)):
        for col in range(0,len(rewardmatrix)):
          # Initialize the sum
          sum = {"up":0,"down":0,"right":0,"left":0}
          if (row,col) == (3,3):
            continue
          for idx,(key,action_value) in enumerate(actions.items()):
            plus = '+' if action_value < len(actions)-1 else ''
            next_row,next_col = next_position(row,col,action_value)
            #Bellman Equation used for each actions
            if is_valid_grid(row=next_row,col=next_col,grid_size=GRID_SIZE):
              customdocument.content += "({t1} *[{t2} + {t3}*{t4}]){p}".format(t1=Trans["curr"],t2=float(rewardmatrix[next_row][next_col]),t3=GAMMA,t4=float(ansmatrix[next_row][next_col]),p=plus)
              sum[key] += Trans["curr"]*(rewardmatrix[next_row][next_col] + GAMMA*ansmatrix[next_row][next_col])
            else:
              customdocument.content +="({t1} *[{t2} + {t3}*{t4}]){p}".format(t1=Trans["other"],t2=float(rewardmatrix[row][col]),t3=GAMMA,t4=float(ansmatrix[row][col]),p=plus)
              sum[key] +=Trans["curr"]*(rewardmatrix[row][col] + GAMMA*ansmatrix[row][col])
              
            
            # Method for Other Actions
            for idx,(next_key,next_action_value) in enumerate(actions.items()):
              if next_key != key:
                plus = '+' if next_action_value < len(actions)-1 else ''
                next_row,next_col = next_position(row,col,next_action_value)
                if is_valid_grid(row=next_row,col=next_col,grid_size=GRID_SIZE):
                  customdocument.content += "({t1} *[{t2} + {t3}*{t4}]){p}".format(t1=Trans["other"],t2=float(rewardmatrix[next_row][next_col]),t3=GAMMA,t4=float(ansmatrix[next_row][next_col]),p=plus)
                  sum[key] += Trans["other"]*(rewardmatrix[next_row][next_col] + GAMMA*ansmatrix[next_row][next_col])
                else:
                  customdocument.content +="({t1} *[{t2} + {t3}*{t4}]){p}".format(t1=Trans["other"],t2=float(rewardmatrix[row][col]),t3=GAMMA,t4=float(ansmatrix[row][col]),p=plus)
                  sum[key] += Trans["other"]*(rewardmatrix[row][col] + GAMMA*ansmatrix[row][col])
            
            customdocument.enwrapactions(key,sum)   
          
          # Receive Q Values of 5 th Iteration          
          if V == 4:
            QArr.append( "\n,".join("{:.3f}".format(e) for e in  list(sum.values())) )
            
            
          maxSum = max(sum.values())
          customdocument.contentmatrix(row,col,maxSum)
        
          print(f"Q Values for Iteration in Row: {row} and column: {col}",list(sum.values()))
          temp_ansmatrix[row][col]=float("{:.3f}".format(maxSum))
          # print(f"")
      head = "V_{no}".format(no=V+1)
      ansmatrix = np.copy(temp_ansmatrix)
      customdocument.addTable(ansmatrix,heading=head)
      customdocument.fullcontent("{head}.docx".format(head=head))


#------------------------------------------------------------------------------------------------------------------------------------------
import numpy as np
GAMMA = 0.9

Trans = {"curr":0.7,"other":0.1,"terminal":1} # Transition Functions where curr is intended state, other is the unintended state , terminal is the terminal state 

"""
# s₁,₁  s₁,₂  s₁,₃  s₁,₄
# s₂,₁  s₂,₂  s₂,₃  s₂,₄
# s₃,₁  *s₃,₂  s₃,₃  *s₃,₄
# s₄,₁  s₄,₂  s₄,₃  s₄,₄
"""
# States defined for Rewards
rewardmatrix = np.full((GRID_SIZE, GRID_SIZE), -0.1)
rewardmatrix[2, 1] = -1  # Hell state s3,2
rewardmatrix[2, 3] = -1  # Hell state s3,4
rewardmatrix[3, 3] = 10

# Initial state of the Answer Matrix
ansmatrix = np.zeros((GRID_SIZE, GRID_SIZE))
ansmatrix[2, 1] = -1  
ansmatrix[2, 3] = -1  
ansmatrix[3, 3] = 10

QArr=[]



prevV = 0

actions= {"up":0,"down":1,"right":2,"left":3}
sum = {"up":0,"down":0,"right":0,"left":0}


valueIterations(GAMMA, Trans, rewardmatrix, ansmatrix, QArr, actions,iterations=1)

# In[8]:



try:
    reshapedArr = np.array(QArr).reshape(4,4)
    Qdocument = CustomDocument()
    headQ = "Q_5"
    Qdocument.addTable(reshapedArr,heading=headQ)
    Qdocument.fullcontent("{head}.docx".format(head=headQ))
except Exception as e:
    print("Exception caused",e)

