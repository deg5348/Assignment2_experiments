#!/usr/bin/env python
# coding: utf-8

# In[1]:



# In[ ]:




# In[2]:


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CustomDocument():
  def __init__(self) -> None:
    self.content=""
    self.mathContentArray = []
    self.contentArray = []
    self.contentaction = ""
    self.doc = Document()
    

  def fullcontent(self,file_name="output.docx",font_name='Times New Roman', font_size=10):
      # Create a new Document
      # Add each line of content to the document
      papra =self.doc.add_paragraph("\n".join(self.contentArray))
      if len(papra.runs) > 0:
        papra_run = papra.runs[0]
        papra_run.font.name = font_name
        papra_run.font.size = Pt(font_size) 

      # Save the document to the specified file
      self.doc.save(file_name)
      self.contentArray=[]

  def contentmatrix(self,row,col,sumresult):
    contentModify = "\n\nV^*_[{row}][{col}] = max({content})  = {res} \n".format(row=row+1,col=col+1,content=self.contentaction,res=sumresult)
    self.contentArray.append(contentModify)
    self.contentaction=""
    
  def addTable(self,data, heading, font_name='Times New Roman', font_size=10):
    heading_paragraph = self.doc.add_heading(level=1)
    heading_run = heading_paragraph.add_run(heading)
    heading_run.font.name = font_name
    heading_run.font.size = Pt(font_size + 2)  
    heading_run.bold = True
    
    # Add a table with the number of rows and columns based on the data
    table = self.doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'  # Optional: Apply a table style
    
    # Populate the table with data and format the text
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            paragraph = cell.add_paragraph(str(cell_text))
            run = paragraph.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = True  # Make the text bold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # self.contentArray.append(table)
    
  def enwrapactions(self,action,sumValue):
    self.contentaction += "\n{act}<{cont} = {sumValue}>\n,".format(act=action,cont=self.content,sumValue=sumValue[action])
    self.content=""
     


# The equation is V<sub>k</sub> = max(up,down,left,right)
# <br>
# up=0,down=1,right=2,left=3
# <br>
# $$
# V^*_s = \ P(s'|a,s) \left[ R_{s,a,s'} + \gamma V^*_{s'} \right]
# $$
# 
# 

# In[3]:


GAMMA = 0.9
REWARD = {"goal": 10, "other": -1}
hellstates = [[2, 1], [2, 3]]
goalstate = [3, 3]
Trans = {"curr": 0.7, "other": 0.1}

rewardmatrix = [
    [-0.1, -0.1, -0.1, -0.1],
    [-0.1, -0.1, -0.1, -0.1],
    [-0.1, -1.0, -0.1, -1.0],
    [-0.1, -0.1, -0.1, 10]
]

ansmatrix = [
    [0, 0, 0, 0],
    [0, 0, 0, 0],
    [0, 0, 0, 0],
    [0, 0, 0, 0]
]

QArr = []

def action_position(row, col, action, rewardmatrix):
    p_row = row
    p_col = col
    if action == 0 and row > 0:  # up
        p_row -= 1
    if action == 1 and row < len(rewardmatrix) - 1:  # down
        p_row += 1
    if action == 2 and col < len(rewardmatrix[0]) - 1:  # right
        p_col += 1
    if action == 3 and col > 0:  # left
        p_col -= 1
    return p_row, p_col

actions = {"up": 0, "down": 1, "right": 2, "left": 3}
sum = {"up": 0, "down": 0, "right": 0, "left": 0}

for V in range(0, 5):
    customdocument = CustomDocument()
    for row in range(0, len(rewardmatrix)):
        for col in range(0, len(rewardmatrix[0])):
            sum = {"up": 0, "down": 0, "right": 0, "left": 0}
            for idx, (key, action_value) in enumerate(actions.items()):
                p_row, p_col = action_position(row, col, action_value, rewardmatrix)
                for val in range(0, len(actions)):
                    if val == action_value:
                        sum[key] += Trans["curr"] * (float(rewardmatrix[p_row][p_col]) + GAMMA * float(ansmatrix[p_row][p_col]))
                    else:
                        sum[key] += Trans["other"] * (float(rewardmatrix[p_row][p_col]) + GAMMA * float(ansmatrix[p_row][p_col]))
                customdocument.enwrapactions(key, sum)
            maxSum = max(sum.values())
            customdocument.contentmatrix(row, col, maxSum)
            ansmatrix[row][col] = float("{:.2f}".format(maxSum))
    head = "V_{no}".format(no=V + 1)
    print(f"Value Matrix for V{V}", ansmatrix)
    customdocument.addTable(ansmatrix, heading=head)
    customdocument.fullcontent("{head}.docx".format(head=head))

# import numpy as np
# reshapedArr = np.array(QArr).reshape(4, 4)
# customdocument = CustomDocument()
# head = "Q_5".format(no=V + 1)
# customdocument.addTable(reshapedArr, heading=head)
# customdocument.fullcontent("{head}.docx".format(head=head))
